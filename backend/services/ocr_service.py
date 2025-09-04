import os
import io
from typing import Optional, Union
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document
from loguru import logger
from config import settings


class OCRService:
    """OCR service for extracting text from various document types"""
    
    def __init__(self):
        """Initialize OCR service"""
        self.tesseract_config = settings.tesseract_config
        self.tesseract_lang = settings.tesseract_lang
        
    async def extract_text(self, file_path: str, mime_type: str) -> Optional[str]:
        """
        Extract text from file based on mime type
        
        Args:
            file_path: Path to the file
            mime_type: MIME type of the file
            
        Returns:
            Extracted text or None if extraction failed
        """
        try:
            if mime_type == "application/pdf":
                return await self._extract_from_pdf(file_path)
            elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                             "application/msword"]:
                return await self._extract_from_docx(file_path)
            elif mime_type.startswith("image/"):
                return await self._extract_from_image(file_path)
            elif mime_type == "text/plain":
                return await self._extract_from_text(file_path)
            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    async def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # First try to extract text directly
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
                
                # If no text found, it might be a scanned PDF - use OCR
                if not text.strip():
                    logger.info(f"No text found in PDF {file_path}, attempting OCR")
                    # For scanned PDFs, you would need to convert to images first
                    # This is a simplified version - in production, you'd use pdf2image
                    return await self._ocr_fallback_for_pdf(file_path)
                    
            return text.strip() if text.strip() else None
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return None
    
    async def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
            return text.strip() if text.strip() else None
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return None
    
    async def _extract_from_image(self, file_path: str) -> Optional[str]:
        """Extract text from image using OCR"""
        try:
            # Open and preprocess image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(
                image, 
                lang=self.tesseract_lang,
                config=self.tesseract_config
            )
            
            return text.strip() if text.strip() else None
            
        except Exception as e:
            logger.error(f"Error extracting text from image {file_path}: {str(e)}")
            return None
    
    async def _extract_from_text(self, file_path: str) -> Optional[str]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text.strip() if text.strip() else None
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return text.strip() if text.strip() else None
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {str(e)}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from text file {file_path}: {str(e)}")
            return None
    
    async def _ocr_fallback_for_pdf(self, file_path: str) -> Optional[str]:
        """OCR fallback for scanned PDFs"""
        try:
            # This is a placeholder for PDF to image conversion + OCR
            # In production, you would use pdf2image library
            logger.info(f"OCR fallback needed for PDF: {file_path}")
            # For now, return None - in real implementation, convert PDF pages to images
            # and run OCR on each page
            return None
            
        except Exception as e:
            logger.error(f"Error in OCR fallback for PDF {file_path}: {str(e)}")
            return None
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword", 
            "text/plain",
            "image/jpeg",
            "image/png",
            "image/tiff",
            "image/bmp"
        ]


# Global OCR service instance
ocr_service = OCRService() 